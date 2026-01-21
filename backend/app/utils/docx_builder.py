from __future__ import annotations

from io import BytesIO

from docx import Document

from app.models import MeetingDocxRequest


def build_meeting_docx(req: MeetingDocxRequest) -> bytes:
    doc = Document()

    title = req.meeting_title or "Meeting Notes"
    doc.add_heading(title, level=1)

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(req.notes.summary or "")

    doc.add_heading("Participants", level=2)
    if req.notes.participants:
        for p in req.notes.participants:
            doc.add_paragraph(p, style="List Bullet")
    else:
        doc.add_paragraph("—")

    doc.add_heading("Conclusions", level=2)
    if req.notes.conclusions:
        for c in req.notes.conclusions:
            doc.add_paragraph(c, style="List Bullet")
    else:
        doc.add_paragraph("—")

    doc.add_heading("Decisions", level=2)
    if req.notes.decisions:
        for d in req.notes.decisions:
            doc.add_paragraph(d, style="List Bullet")
    else:
        doc.add_paragraph("—")

    doc.add_heading("Action Items", level=2)
    if req.notes.action_items:
        for a in req.notes.action_items:
            owner = f" ({a.owner})" if a.owner else ""
            due = f" [due: {a.due_date}]" if a.due_date else ""
            doc.add_paragraph(f"{a.item}{owner}{due}", style="List Bullet")
    else:
        doc.add_paragraph("—")

    doc.add_heading("Transcript", level=2)
    doc.add_paragraph(req.transcript or "")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

